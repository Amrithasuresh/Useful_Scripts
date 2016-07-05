# -------------------------------------------------------------------------------
# Purpose:     How to create a word file using python?
# Author:      Suresh Panneerselvam
#
# Created:     2016-07-05
# OS:          Ubuntu 16.04
# Licence:     MIT
# -------------------------------------------------------------------------------
#import the docx module
#Documentation can be found here http://python-docx.readthedocs.io/en/latest/
import docx

document = docx.Document()

p = document.add_paragraph('Sequence Patterns')

table = document.add_table(rows=1, cols=5)
hdr_cells = table.rows[0].cells
hdr_cells[0].text =  'S. No'
hdr_cells[1].text = 'Gene ID'
hdr_cells[2].text = 'Sequence Length'
hdr_cells[3].text = 'Start-End'
hdr_cells[4].text = 'Strand Type'

recordset = [
    {'S.No': 1, 'Gene ID': "NM_172887.2", 'Sequence Length': '1-10753','Start-End': '10453-10459','Strand Type':'+'},
    {'S.No': 2, 'Gene ID': "NM_172887.2", 'Sequence Length': '1-10753','Start-End': '10453-10459','Strand Type':'+'},
    {'S.No': 3, 'Gene ID': "NM_172887.2", 'Sequence Length': '1-10753','Start-End': '10453-10459','Strand Type':'+'},
    {'S.No': 4, 'Gene ID': "NM_172887.2", 'Sequence Length': '1-10753','Start-End': '10453-10459','Strand Type':'+'},
    {'S.No': 5, 'Gene ID': "NM_172887.2", 'Sequence Length': '1-10753','Start-End': '10453-10459','Strand Type':'+'},

]

for item in recordset:
    row_cells = table.add_row().cells
    row_cells[0].text = str(item['S.No'])
    row_cells[1].text = str(item['Gene ID'])
    row_cells[2].text = str(item['Sequence Length'])
    row_cells[3].text = str(item['Start-End'])
    row_cells[4].text = str(item['Strand Type'])

document.save('demo.docx')